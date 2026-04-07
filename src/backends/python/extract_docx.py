#!/usr/bin/env python3
"""Extract Markdown text + section map + tables from a .docx file."""
import sys, json
from docx.document import Document as DocxDocument
from docx.table import Table
from docx.text.paragraph import Paragraph

def main():
    if len(sys.argv) < 2:
        print(json.dumps({"error": "Usage: extract_docx.py <filepath>"}))
        sys.exit(1)
    filepath = sys.argv[1]
    try:
        from docx import Document
    except ImportError:
        print(json.dumps({"error": "python-docx not installed. Run: pip install python-docx"}))
        sys.exit(1)
    try:
        doc: DocxDocument = Document(filepath)
    except Exception as e:
        print(json.dumps({"error": f"Failed to open Docx: {e}"}))
        sys.exit(1)

    # First pass: collect all document elements (paragraphs and tables) in order
    # This is needed because tables can appear between paragraphs
    from collections import defaultdict
    section_tables = defaultdict(list)  # section_idx -> list of tables

    sections = []
    lines = []
    line_num = 1
    current = None
    section_idx = 0

    # Get all elements (paragraphs and tables) in document order
    # We'll iterate through the document body which contains both paragraphs and tables
    for element in doc.element.body:
        # Check if this element is a paragraph
        if element.tag.endswith('p'):
            para = None
            for p in doc.paragraphs:
                if p._element == element:
                    para = p
                    break
            if para is None:
                continue

            style = para.style.name
            text = para.text
            if style.startswith("Heading"):
                try:
                    level = int(style.split()[-1])
                except ValueError:
                    level = 1
                if current:
                    current["line_end"] = line_num - 1
                    sections.append(current)
                current = {"section_idx": section_idx, "heading": text.strip(), "level": level, "line_start": line_num, "line_end": None}
                section_idx += 1
                lines.append("#" * level + " " + text.strip())
            else:
                lines.append(text)
            line_num += 1

        # Check if this element is a table
        elif element.tag.endswith('tbl'):
            table = None
            for t in doc.tables:
                if t._element == element:
                    table = t
                    break
            if table is None:
                continue

            rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            html = "<table>" + "".join("<tr>" + "".join(f"<td>{c}</td>" for c in row) + "</tr>" for row in rows) + "</table>"
            # Associate table with current section (or section 0 if no sections yet)
            current_section_idx = current["section_idx"] if current else 0
            section_tables[current_section_idx].append(html)

    # Close the last section
    if current:
        current["line_end"] = line_num - 1
        sections.append(current)

    # If no headings, create a single implicit section
    if not sections:
        sections.append({"section_idx": 0, "heading": "", "level": 1, "line_start": 1, "line_end": line_num - 1})

    # Build tables list with proper section_idx association
    tables = []
    for sec_idx, html_list in sorted(section_tables.items()):
        for html in html_list:
            tables.append({"section_idx": sec_idx, "html": html})

    print(json.dumps({"markdown": "\n".join(lines), "sections": sections, "tables": tables}))

if __name__ == "__main__":
    main()
